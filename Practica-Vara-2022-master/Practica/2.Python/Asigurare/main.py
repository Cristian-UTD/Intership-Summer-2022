
import datetime
import docx
from datetime import date
from datetime import datetime
from itp_asigurare import cars

masini_itp = cars.itp

masini_asigurare = cars.asigurare

ziua_astazi = date.today()

ziua_astazi_ora = datetime.today()

masini_expirate_asigurare = {}

masini_amenda_asigurare = {}

masini_expirate_itp = {}

masini_amenda_itp = {}

valabila_asigurare = {}

valabila_itp = {}

masini = []

def verificare_itp():
    global valabila_itp
    data_itp = datetime.strptime(masini_itp[masina], '%d.%m.%Y')
    data_itp = data_itp.date()
    if (data_itp > ziua_astazi):
        valabila_itp[f"{masina}"] = "Ok"
    else:
        data_itp = data_itp.strftime('%d.%m.%Y')
        masini_expirate_itp[f"{masina}"] = data_itp
        amenda_itp()
        valabila_itp[f"{masina}"] = "NOk"

def verificare_asigurare():
    global valabila_asigurare
    data_asigurare = datetime.strptime(masini_asigurare[masina], '%d.%m.%Y')
    data_asigurare = data_asigurare.date()
    if (data_asigurare > ziua_astazi):
        valabila_asigurare[f"{masina}"] = "Ok"
    else:
        data_asigurare = data_asigurare.strftime('%d.%m.%Y')
        masini_expirate_asigurare[f"{masina}"] = data_asigurare
        amenda_asigurare()
        valabila_asigurare[f"{masina}"] = "NOk"

def amenda_itp():
    data_itp = datetime.strptime(masini_itp[masina], '%d.%m.%Y')
    data_itp = data_itp.date()
    diferenta = ziua_astazi - data_itp
    diferenta = diferenta.days
    if (diferenta > 30):
        if (diferenta > 90):
            if(diferenta > 365):
                masini_amenda_itp[f"{masina}"] = 5000
            else:
                masini_amenda_itp[f"{masina}"] = 2000
        else:
            masini_amenda_itp[f"{masina}"] = 1000
    else:
        masini_amenda_itp[f"{masina}"] = 300

def amenda_asigurare():
    data_asigurare = datetime.strptime(masini_asigurare[masina], '%d.%m.%Y')
    data_asigurare = data_asigurare.date()
    diferenta = ziua_astazi - data_asigurare
    diferenta = diferenta.days
    if (diferenta > 30):
        if (diferenta > 90):
            if(diferenta > 365):
                masini_amenda_asigurare[f"{masina}"] = 5000
            else:
                masini_amenda_asigurare[f"{masina}"] = 2000
        else:
            masini_amenda_asigurare[f"{masina}"] = 1000
    else:
        masini_amenda_asigurare[f"{masina}"] = 300

def logs(masinile):
    global doc
    doc = docx.Document()
    for x in masinile:
        if (valabila_itp[x] == "Ok" and valabila_asigurare[x] == "Ok"):
            doc.add_paragraph("____________________________________________________________________________________")
            doc.add_paragraph(f"[{ziua_astazi_ora}] [{x}] Asigurare {valabila_asigurare[x]}/ ITP {valabila_itp[x]}")
        else:
            if (valabila_itp[x] == "NOk" and valabila_asigurare[x] == "NOk"):
                doc.add_paragraph("____________________________________________________________________________________")
                doc.add_paragraph(f"[{ziua_astazi_ora}] [{x}] Asigurare {valabila_asigurare[x]}/ ITP {valabila_itp[x]} ({masini_amenda_itp[x] + masini_amenda_asigurare[x]} RON)")
            else:
                if (valabila_itp[x] == "Ok" and valabila_asigurare[x] == "NOk"):
                    doc.add_paragraph("____________________________________________________________________________________")
                    doc.add_paragraph(f"[{ziua_astazi_ora}] [{x}] Asigurare {valabila_asigurare[x]}/ ITP {valabila_itp[x]} ({masini_amenda_asigurare[x]} RON)")
                else:
                    if (valabila_itp[x] == "NOk" and valabila_asigurare[x] == "Ok"):
                        doc.add_paragraph("____________________________________________________________________________________")
                        doc.add_paragraph(f"[{ziua_astazi_ora}] [{x}] Asigurare {valabila_asigurare[x]}/ ITP {valabila_itp[x]} ({masini_amenda_itp[x]} RON)")

    doc.save("Logs.docx")

def verificare():
    global masina
    for masina in masini_itp:
        verificare_itp()
    for masina in masini_asigurare:
        verificare_asigurare()
    for masina in masini_itp:
        masini.append(masina)
    logs(masini)

class Verificare():
    verificare()

print("Masini expirate_itp:")
print(masini_expirate_itp)
print("Amenzi:")
print(masini_amenda_itp)

print("Masini expirate_asigurare:")
print(masini_expirate_asigurare)
print("Amenzi:")
print(masini_amenda_asigurare)
